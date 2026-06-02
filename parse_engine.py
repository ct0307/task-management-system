"""
本地文件解析引擎
支持: PDF(文字+扫描件OCR) / 图片OCR / Word(.docx)
用法: python parse_engine.py <filepath>
"""
import sys, os, json

def _check_deps(names, install_hint):
    for name in names:
        try: __import__(name)
        except ImportError: return f"需要安装: pip install {' '.join(names)} ({install_hint})"
    return None

def parse_pdf_text(filepath):
    """提取 PDF 文字层"""
    err = _check_deps(['pdfplumber'], 'pip install pdfplumber')
    if err: return {'error': err}
    import pdfplumber
    pages_text = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages_text.append({'page': i+1, 'text': text})
    if not pages_text:
        return {'error': 'PDF 无可提取文字', 'pages': []}
    return {
        'pages': pages_text,
        'full_text': '\n'.join(p['text'] for p in pages_text)
    }

def parse_pdf_ocr(filepath):
    """扫描件 PDF → 图片 → OCR"""
    err = _check_deps(['pdf2image', 'PIL', 'pytesseract'], 'pip install pdf2image Pillow pytesseract')
    if err: return {'error': err}
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    try:
        images = convert_from_path(filepath, dpi=200, first_page=1, last_page=10)
    except Exception as e:
        try_images = convert_from_path(filepath, dpi=150, first_page=1, last_page=5, fmt='jpeg')
        images = try_images
    pages_text = []
    for i, img in enumerate(images):
        text = pytesseract.image_to_string(img, lang='chi_sim+eng')
        if text.strip():
            pages_text.append({'page': i+1, 'text': text})
    if not pages_text:
        return {'error': 'OCR 未识别到文字'}
    return {
        'pages': pages_text,
        'full_text': '\n'.join(p['text'] for p in pages_text)
    }

def parse_pdf(filepath):
    """智能 PDF 解析：先试课表 → 文字 → OCR"""
    try:
        from schedule_parser import parse as parse_schedule
        result = parse_schedule(filepath)
        if len(result) >= 3:
            return {
                'type': 'schedule',
                'courses': result,
                'summary': '\n'.join(
                    f"{c['day']} P{c['period']} {c['start_time']}-{c['end_time']} {c['name']}"
                    + (f" 地点:{c['location']}" if c.get('location') else '')
                    + (f" 教师:{c['teacher']}" if c.get('teacher') else '')
                    for c in result
                )
            }
    except Exception: pass

    result = parse_pdf_text(filepath)
    if 'error' in result and '无可提取' in result.get('error', ''):
        return parse_pdf_ocr(filepath)
    return result

def parse_docx(filepath):
    """Word 文档解析"""
    err = _check_deps(['docx'], 'pip install python-docx')
    if err: return {'error': err}
    from docx import Document
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(rows)
    return {
        'paragraphs': paragraphs,
        'tables': tables,
        'full_text': '\n'.join(paragraphs)
    }

def parse(filepath):
    """自动识别文件类型并解析"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return parse_pdf(filepath)
    elif ext in ('.png', '.jpg', '.jpeg', '.webp', '.bmp'):
        return ocr_image(filepath)
    elif ext == '.docx':
        return parse_docx(filepath)
    else:
        return {'error': f'不支持的文件格式: {ext}，支持 .pdf .png .jpg .docx'}

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({'error': '用法: python parse_engine.py <filepath>'}, ensure_ascii=False))
        sys.exit(1)
    result = parse(sys.argv[1])
    print(json.dumps(result, ensure_ascii=False, indent=2))
